from typing import Dict, Any
from .interface import OutputGeneratorInterface

try:
    from docx import Document
except ImportError:
    Document = None

class DOCXOutputGenerator(OutputGeneratorInterface):
    def generate(self, content: Any, filename: str, metadata: Dict[str, Any] = None) -> Dict[str, Any]:
        if Document is None:
            return {
                "status": "failed",
                "format": "docx",
                "errors": ["python-docx is not installed."],
                "metadata": metadata or {}
            }
            
        if not filename.endswith('.docx'):
            filename += '.docx'
            
        safe_path = self._get_safe_path(filename)
        metadata = metadata or {}
        
        try:
            doc = Document()
            # If content is a list of strings, add them as paragraphs
            if isinstance(content, list):
                for item in content:
                    doc.add_paragraph(str(item))
            elif isinstance(content, dict):
                for k, v in content.items():
                    doc.add_heading(str(k), level=1)
                    doc.add_paragraph(str(v))
            else:
                doc.add_paragraph(str(content))
                
            doc.save(safe_path)
                
            return {
                "status": "success",
                "format": "docx",
                "path": safe_path,
                "filename": filename,
                "metadata": metadata
            }
        except Exception as e:
            return {
                "status": "failed",
                "format": "docx",
                "errors": [str(e)],
                "metadata": metadata
            }
